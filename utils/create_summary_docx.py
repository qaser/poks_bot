from math import ceil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from config.mongo_config import petitions
import utils.constants as const


def set_cell_color(cell, color):
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color)
    table_cell_properties.append(shade_obj)


async def create_docx_file():
    DOC_HEADER = 'Отчет по проблемным вопросам филиалов по направлению ПОпоЭКС'
    TABLE_HEADERS = (
        '№ п/п',
        'Направление',
        'Дата',
        'Содержание вопроса',
        'Статус',
    )
    document = Document()
    section = document.sections[0]
    # левое поле в миллиметрах
    section.left_margin = Mm(25)
    # правое поле в миллиметрах
    section.right_margin = Mm(15)
    # верхнее поле в миллиметрах
    section.top_margin = Mm(15)
    # нижнее поле в миллиметрах
    section.bottom_margin = Mm(15)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    header = document.add_paragraph(DOC_HEADER)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.style = document.styles['Normal']

    table = document.add_table(rows=1, cols=5, style='Table Grid')
    table_hdr = table.rows[0].cells  # заголовки таблицы
    table.columns[0].width = Mm(10)
    table.columns[1].width = Mm(35)
    table.columns[2].width = Mm(25)
    table.columns[3].width = Mm(80)
    table.columns[4].width = Mm(28)

    for id, cell in enumerate(table_hdr):
        cell.text = TABLE_HEADERS[id]

    current_rows = 1  # текущее количество строк в таблице
    for ks in const.KS:
        queryset = list(petitions.find(
            {'ks': ks, 'status': {'$in': ['create', 'inwork', 'finish']}}
        ).sort([('status', 1), ('directions', 1)]))
        num_petitions = len(queryset)
        if num_petitions == 0:
            continue
        ks_row = table.add_row().cells
        ks_row[0].merge(ks_row[4])
        ks_row[0].text = ks
        ks_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_color(ks_row[0], 'B5B8B1')

        for num, pet in enumerate(queryset):
            status = pet.get('status')
            main_row = table.add_row().cells
            main_row[0].text = str(num + 1)
            main_row[1].text = const.DIRECTIONS_CODES[pet.get('direction')]
            main_row[2].text = pet.get('date')
            main_row[3].text = pet.get('text')
            main_row[4].text = const.PETITION_STATUS[status][0]
            set_cell_color(main_row[4], const.PETITION_COLOR[status])
            current_rows += 1

    path = f'static/docs_email/Сводный перечень вопросов.docx'
    document.save(path)
